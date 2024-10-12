from docx import Document
import os


def produce_report(data: dict,
                   template_path: str = './static/docx/report_format.docx',
                   output_path: str = './outputs/docx',
                   file_name: str = 'report.docx',
                   supress_errors: bool = False) -> None:
    """Produces a report of the data to be stored in a certain relative
    location based on a template.

    Args:
        data (dict): The data to be replaced in the docx template.
        template_path (str): The location of the template to be used.
            default: "./static/docx/report_format.docx"
        output_path (str): The output path of the processed docx file.
            default: r"./outputs/docx"
        file_name (str): The file name to be saved as.
            default: r"report.docx"
        supress_errors (bool): If the FileExistsError must be supressed.
            default: False

    Returns: No return value.
    """
    if os.path.exists(output_path+'/'+file_name):
        if not supress_errors:
            raise FileExistsError(
                'The given output path already has an existing file.')
        else:
            os.remove(output_path+'/'+file_name)

    if not os.path.exists(output_path):
        os.makedirs(output_path)

    template = Document(template_path)
    set_variables(data, template)
    set_tables(data, template)

    template.save(output_path+'/'+file_name)


def set_variables(data: dict, template: Document) -> None:
    """Replaces the keys in the given Document object with the values
    provided in the dictionary.

    Args:
        data (dict): Values and keys.
        template (docx.Document): The docx template to replace in.

    Returns: No return value.
    """

    for i in data:
        if not i.startswith('$T'):
            for paragraph in template.paragraphs:
                if paragraph.text.find(i) >= 0:
                    paragraph.text = paragraph.text.replace(i, data[i])

    for key, value in data.items():
        if not key.startswith('$T'):
            for table in template.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph.text = paragraph.text.replace(key, value)


def set_tables(data: dict, template: Document):
    for i in data.keys():
        if '$T' in i:
            for paragraph in template.paragraphs:
                if paragraph.text.find(i) >= 0:
                    paragraph.text = ''

                    table_data = data[i]
                    # the number of rows of data, including the header row
                    rows_no = len(table_data[[*table_data][0]]) + 1

                    table = template.add_table(
                        rows=rows_no, cols=len(table_data))
                    table_header = table.rows[0].cells

                    # to create the header of the table. Separate from the main loop
                    # for the rest of the data so that if formatting is needed,
                    # it is easy to do.
                    for col in range(len(table_data)):
                        table_header[col].text = str([*table_data][col])

                    for row in range(1, rows_no):
                        table_row = table.rows[row].cells
                        for col in range(len(table_data)):
                            table_row[col].text = str(
                                table_data[[*table_data][col]][row - 1])
                            # 'row - 1' since the first row is taken up by the header

                    # to place the table in place of the variable in the template.
                    location = paragraph._element.getparent()
                    location.insert(location.index(
                        paragraph._element), table._element)


# FOR TESTING PURPOSES
if __name__ == '__main__':
    produce_report({
        '${ADMN_NO}': 'A1455', '${STUDENT_NAME}': 'John Doe',
        '$T{SCHOLASTIC_AREAS_TABLE}':
            {
                'Subjects': ['maths', 'Computer Science', 'phys'],
                'Marks': [23, 34, 90],
                'Grade': [None, None, None]
            }
    }, supress_errors=True)
